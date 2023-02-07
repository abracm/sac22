import csv
import docx

dados = []

with open("resources\\competitors-and-staff (2).csv", "r", encoding="utf-8") as f:
    linhas = csv.DictReader(f, delimiter=';')
    for row in linhas:
        dados.append(row)

traducao = {"Argentina": "Argentina",
            "Australia":"Austrália",
            "Bolivia": "Bolivia",
            "Brazil": "Brasil",
            "Canada": "Canadá",
            "Chile":"Chile",
            "Colombia":"Colombia",
            "India":"Índia",
            "USA":"Estados Unidos",
            "United States":"Estados Unidos",
            "France":"França",
            "Panama":"Panamá",
            "Syria":"Síria",
            "United Kingdom":"Reino Unido"}

for row in dados:
    row["PaisesPtbrEs"] = traducao[row["Country"]]

dados = sorted(dados, key=lambda d: d['Name'])

document = docx.Document("resources\\name-tags-template.docx")
table = document.tables[0]
linha_atual = 0
coluna_atual = 0
linha_da_pagina = 0
for row in dados:
    cell = table.cell(linha_atual, coluna_atual)
   
    if linha_da_pagina in (3,4,5,6):
        
        cell.add_paragraph().style = ("Micro")        
        paragraph = cell.paragraphs[0]
        p = paragraph._element  
        p.getparent().remove(p)
        p._p = p._element = None
        
        cell.add_paragraph().text = row["Name"]
    
    else:
        cell.text = row["Name"]

    if len(row["Name"]) > 28:
        cell.paragraphs[-1].style = "NomeLongo"
    elif len(row["Name"]) > 24:
        cell.paragraphs[-1].style = "NomeMedio"
    else:
        cell.paragraphs[-1].style = "Nome"
    cell.add_paragraph().style = "País"
    paragraph = cell.paragraphs[-1]
    if row["WCA ID"]!= "":
        r = paragraph.add_run()    
        r.text = row["WCA ID"] + " - "
    
    r = paragraph.add_run()
    r.text = row["PaisesPtbrEs"] + " "
    r = paragraph.add_run()
    r.add_picture("resources\\countries\\" + row["Country"] + ".png", 
                  width = docx.shared.Cm(1))
    r.font.subscript = True
    
    if linha_da_pagina == 6 and coluna_atual == 1:
        linha_da_pagina = 0
    elif coluna_atual == 1:
        linha_da_pagina += 1
    
    if coluna_atual == 0:
        coluna_atual = 1
    else:
        coluna_atual = 0
        linha_atual +=1
document.save("name-tags.docx")